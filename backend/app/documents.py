"""Extrai texto de PDF/Word para a base da empresa (com OCR em PDF escaneado)."""
from __future__ import annotations

from io import BytesIO

MAX_UPLOAD_BYTES = 5 * 1024 * 1024
MAX_PDF_PAGES = 80
MAX_OCR_PAGES = 15
MIN_TEXT_CHARS = 20


def extract_text_from_upload(filename: str, data: bytes) -> tuple[str, str, dict]:
    """
    Retorna (conteúdo, source_type, meta).
    meta pode incluir ocr=True quando usou reconhecimento de imagem.
    """
    if len(data) > MAX_UPLOAD_BYTES:
        raise ValueError("Arquivo acima de 5 MB. Envie um PDF/Word menor.")
    name = (filename or "").lower().strip()
    if name.endswith(".pdf"):
        content, meta = _from_pdf(data)
        return content, "pdf", meta
    if name.endswith(".docx"):
        return _from_docx(data), "docx", {"ocr": False}
    if name.endswith(".doc"):
        raise ValueError("Use o formato .docx (Word moderno), não .doc antigo.")
    raise ValueError("Formato não suportado. Envie PDF (.pdf) ou Word (.docx).")


def _from_pdf(data: bytes) -> tuple[str, dict]:
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(data))
    page_count = min(len(reader.pages), MAX_PDF_PAGES)
    parts: list[str] = []
    for page in reader.pages[:MAX_PDF_PAGES]:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text.strip())
    content = "\n\n".join(parts).strip()
    if len(content) >= MIN_TEXT_CHARS:
        return content, {"ocr": False, "pages": page_count, "method": "text"}

    # PDF sem texto selecionável → OCR (até MAX_OCR_PAGES)
    try:
        ocr_content, ocr_pages = _ocr_pdf(data)
    except Exception as exc:
        raise ValueError(
            "Não deu para ler este PDF (parece escaneado) e o OCR falhou. "
            f"Detalhe: {exc}. Tente um PDF com texto selecionável ou cole o conteúdo."
        ) from exc
    if len(ocr_content) < MIN_TEXT_CHARS:
        raise ValueError(
            "OCR não encontrou texto legível neste PDF. "
            "Verifique se o scan está nítido ou cole o conteúdo manualmente."
        )
    return ocr_content, {
        "ocr": True,
        "pages": ocr_pages,
        "method": "ocr",
        "note": f"Lido por OCR em {ocr_pages} página(s) (limite {MAX_OCR_PAGES}).",
    }


def _ocr_pdf(data: bytes) -> tuple[str, int]:
    import fitz  # pymupdf
    import pytesseract
    from PIL import Image

    doc = fitz.open(stream=data, filetype="pdf")
    parts: list[str] = []
    pages_done = 0
    for i, page in enumerate(doc):
        if i >= MAX_OCR_PAGES:
            break
        # 2x melhora OCR sem estourar memória em páginas A4
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
        img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
        try:
            text = pytesseract.image_to_string(img, lang="por+eng")
        except pytesseract.TesseractError:
            text = pytesseract.image_to_string(img, lang="eng")
        pages_done += 1
        if text and text.strip():
            parts.append(text.strip())
    doc.close()
    return "\n\n".join(parts).strip(), pages_done


def _from_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(BytesIO(data))
    parts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    content = "\n".join(parts).strip()
    if len(content) < MIN_TEXT_CHARS:
        raise ValueError("Documento Word sem texto suficiente para publicar.")
    return content
